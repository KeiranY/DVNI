import random

from docx import Document

from container.kali import Kali
from scenarios import AccessVlan, Scenario

# TODO: Add a second kali machine in another VLAN in another switch, allow the user to use that one to receive packets

# Explain how double tagging works and why it is unidirectional

# What vlan is the second kali machine on? (2)
# Use vconfig on the first kali machine to add it's own vlan to it's interface, a new interface with the name kali-eth0.8 should be added
# Use vconfig on the created interface to add the second kali's vlan, a new interface with the name kali-eth0.8.2 should be added
# Use this interface to arp-scan
# Are any devices found? (No)
# Use wireshark/tcpdump on the second Kali, and run the scan again
# Is any traffic received? (Yes)
# What Vlan, if any, does the traffic belong to? (None)
# Is a response sent? (Yes)
# What Vlan is the response sent to? (None)
from utils.vlan import VlanMode


class Import(AccessVlan.Import):
    name = "VLAN Double Tagging"
    enabled = True
    weight = 70

    kali_receive = None
    receive_mode = None

    def create_network(self, controller=None):
        super(Import, self).create_network(controller)
        # Add a trunk to Kali's switch port to only accept packets with
        self.switches[0].addTrunk(self.kali.defaultIntf().link.intf1, self.vlans[0])
        # Add a second Kali to the second switch in the second VLAN
        self.kali_receive = self.net.addHost('kali2',
                                             cls=Kali,
                                             vnc=Kali.VNC_DEFAULT + 1,
                                             web=Kali.WEB_DEFAULT + 1,
                                             ip="%s/%s" % (self.hosts.pop(), self.prefixlen))
        link = self.net.addLink(self.switches[1], self.kali_receive)
        self.switches[1].addTag(link.intf1, self.vlans[1])
        # If the mode is ACCESS the Kali machine will receive packets with no VLAN tag
        self.receive_mode = random.choice([VlanMode.ACCESS, VlanMode.TAGGED])
        self.switches[1].addMode(link.intf1, self.receive_mode)

    def run_network(self):
        super(Import, self).run_network()
        defaultintf = str(self.kali.defaultIntf())
        self.net.staticArp()
        if self.developer:
            self.kali.cmd("vconfig add %s %d" % (defaultintf, self.vlans[0]))
            nativeintf = defaultintf + '.' + str(self.vlans[0])
            self.kali.cmd("ip l set up dev %s" % nativeintf)
            for v in self.vlans:
                self.kali.cmd("vconfig add %s %d" % (nativeintf, v))
                doubletagintf = nativeintf + '.' + str(v)
                self.kali.cmd("ip l set up dev %s" % doubletagintf)
                self.kali.cmd("arp-scan -I %s %s/%s" % (doubletagintf, self.kali.defaultIntf().IP(), self.prefixlen))

    def generate_task(self, doc):
        Scenario.generate_task(self, doc)
        doc.add_paragraph("In this scenario you have access to two Kali machines in seperate VLANs in the same network. "
                          "The VLANs are somewhere in the range 1-10. ")
        doc.add_paragraph(
            "The tool arp-scan is able to which VLAN you are a part of by specifying a VLAN to scan.")
        doc.add_paragraph("Answer the following questions: ")

    def generate_questions(self):
        self.questions += [("What VLAN is the first Kali machine connected to", str(self.vlans[0])),
                           ("What VLAN is the second Kali machine connected to", str(self.vlans[1])),
                           ("Use vconfig on the first Kali machine to add a VLAN tag to it's task interface. What is the new interface's name", "%s.%s" % (self.kali.defaultIntf(), self.vlans[0])),
                           ("Repeat this process to add a VLAN tag to the new interface, this time adding the second Kali machine's VLAN. What is the new interface's name", "%s.%s.%s" % (self.kali.defaultIntf(), self.vlans[0], self.vlans[1])),
                           ("Use this interface to arp-scan for devices, are any devices found", "No, as VLAN tagging is unidirectional"),
                           ("Open wireshark on the second Kali and run the scan again, was any traffic received","Yes"),
                           ("If so, what VLAN, if any, does it belong to", "None" if self.receive_mode is VlanMode.ACCESS else str(self.vlans[1])),
                           ("Was any traffic sent", "Yes" if self.receive_mode is VlanMode.ACCESS else "No"),
                           ("If so, what VLAN, if any, does it belong to", "None" if self.receive_mode is VlanMode.ACCESS else "N/A")]


if __name__ == "__main__":
    # Run network
    Import(developer=True, seed="debug").run()
