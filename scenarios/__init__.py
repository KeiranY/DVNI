"""
Scenarios
=========
"""

import logging
import os
import random

from docx import Document
from mininet.clean import cleanup, killprocs, Cleanup
from mininet.cli import CLI
from mininet.log import info, setLogLevel
from mininet.net import Containernet
from mininet.node import Controller
from pyftpdlib.authorizers import DummyAuthorizer
from pyftpdlib.handlers import FTPHandler
from pyftpdlib.servers import MultiprocessFTPServer

from container.kali import Kali
from controller import PoxController

from typing import List, Tuple

# Add a cleanup command to mininet.clean to clean pox controller
Cleanup.addCleanupCallback(lambda: killprocs(PoxController.pox_comand))


class Scenario(object):
    """
    Base scenario class.
    Performs the functions to run and document a scenario.
    Should be extended by any other implemented scenarios, all public methods are extensible.
    Extension classes should have the class name `Import` as per the `start` file.
    """
    # These attributes allow for the filtering/ordering of scenarios when presenting them to students
    name = "Base Scenario"
    """Scenario name. Used in the produced documentation"""
    enabled = False
    """If the scenario should be shown to users."""
    weight = -1
    """Used to order scenarios before presenting them to users, the lower the weight the earlier in the list."""

    def __init__(self, teacher=False, developer=False, seed=None):
        # type: (bool, bool, str) -> None
        """
        :type teacher: bool
        :param teacher: Is the user a teacher.
        :type developer: bool
        :param developer: Is the user a developer.
        :type seed: str
        :param seed: str: Random number generator seed, likely student ID.
        """
        self.teacher = teacher
        self.developer = developer
        self.net = None  # type: Containernet
        """Containernet: The network of devices."""
        self.questions = []  # type: List[Tuple[str, str]]
        """List[Tuple[str, str]]: List of question & answer combinations used for task & answer documents."""
        self.answer_document = Document()
        """Document: `python-docx` document used to provide the scenario answers to teachers."""
        self.task_document = Document()
        """Document: `python-docx` document used to provide the scenario task to users."""
        # Use the inputted seed, if given, to randomise the network
        self.seed = seed
        """Seed parameter passed to __init__, used to randomise the network parameters. 
        Added to answer sheets given to teachers so they know the student's ID."""
        if self.seed is not None:
            random.seed(self.name + self.seed)
            # If no seed, set seed to "random"
            # to indicate the outputted docuemnts are for a random seed
        else:
            self.seed = "random"

    def run(self):
        """Main function that executes the rest of the scenario."""
        # Default to only error log printing for students
        setLogLevel('error')
        if self.teacher:
            # Allow teaches to see more complete logs
            setLogLevel('info')
        if self.developer:
            # Show debug logs for development
            setLogLevel('debug')
        # Create Containernet network
        self.create_network()
        # Start Containernet network
        self.run_network()
        # Generate task/answer sheets
        self.generate_task(self.task_document)
        self.generate_questions()
        self._add_questions(self.task_document)
        self._add_answers(self.answer_document)
        self.save_documents()
        if self.developer:
            # If we're a developer, start the CLI
            # so we can test from the command line
            CLI(self.net)
            self.net.stop()
        return

    def create_network(self, controller=Controller):
        """Create Containernet network."""
        info('*** Running Cleanup\n')
        cleanup()
        self.net = Containernet(controller=controller)
        if controller is not None:
            self.add_controller()

    def add_controller(self):
        """Adds a controller to the network."""
        self.net.addController()

    def _add_answers(self, doc):
        doc.add_heading('Answers', level=3)
        for idx, question in enumerate(self.questions):
            # If the question has 2 items print it as a question and answer
            if question[1] != "":
                doc.add_paragraph(str(idx + 1) + '. ' + question[0] + ': ').add_run(question[1]).bold = True
            # Else don't print it

    def generate_task(self, doc):
        """Adds a header and information on how to connect to the network (if applicable).
        Scenarios should extend this to provide users with details on the task."""
        doc.add_heading('%s' % self.name, level=2)
        kali = self.net.get('kali')  # type: Kali
        if kali:
            kali.add_hint(doc)

    def generate_questions(self):
        pass

    def _add_questions(self, doc):
        doc.add_heading('Questions', level=3)
        for idx, question in enumerate(self.questions):
            # If the question has 2 items print it as a question
            if question[1] != "":
                doc.add_paragraph(str(idx + 1) + '. ' + question[0] + '? ')
            # Else print it as a statement
            else:
                doc.add_paragraph(question[0])

    def save_documents(self, studentDirectory='./student/', teacherDirectory='./teacher/', studentAllowedAnswers=False,
                       staticTaskDocument=True):
        """
        Saves the task and answer documents locally so they can be provided to users.

        :param studentDirectory: Location of the student-accessible folder on the VM
        :type studentDirectory: string

        :param teacherDirectory: Location of the teacher-accessible folder on the VM
        :type teacherDirectory: string

        :param studentAllowedAnswers: Is the student provided the answer document for the task.
        :type studentAllowedAnswers: bool

        :param staticTaskDocument: Is the task document the same for every student.
        :type staticTaskDocument: bool
        """
        # Create neccisary directories if they do not exist
        for directory in [studentDirectory, teacherDirectory]:
            if not os.path.exists(directory):
                os.makedirs(directory)

        # -- Answer sheet --

        # For teachers, generate a document named "[Scenario name]-[Stduent ID].docx"
        self.answer_document.save(teacherDirectory + self.name + '-' + self.seed + '.docx')
        # If we're a student, generate a document named "[Scenario name]-answers.docx"
        if studentAllowedAnswers:
            self.answer_document.save(studentDirectory + self.name + '-answers.docx')

        # -- Task sheet --

        # If the task sheet is the same for all students
        if staticTaskDocument:
            taskLocation = teacherDirectory + self.name + '.docx'
            # Create the task sheet
            self.task_document.save(taskLocation)
        # If we're a student
        # if not self.teacher:
        self.task_document.save(studentDirectory + self.name + '.docx')

    def run_ftp(self, studentDirectory='./student/', teacherDirectory='./teacher'):
        """
        Runs a PyFTPlib FTP server to provide users with access to their task/answer documents.

        :param studentDirectory: Location of the student-accessible folder on the VM
        :type studentDirectory: string

        :param teacherDirectory: Location of the teacher-accessible folder on the VM
        :type teacherDirectory: string
        """

        authorizer = DummyAuthorizer()
        pw = '%05x' % random.randrange(16 ** 5)
        if self.teacher:
            authorizer.add_user(username="teacher", password=pw, homedir=teacherDirectory)
        else:
            authorizer.add_user(username="student", password=pw, homedir=studentDirectory)
        handler = FTPHandler
        handler.authorizer = authorizer

        handler.banner = "DVNG tasks ftp server"

        # Instantiate FTP server class and listen on 0.0.0.0:2121
        address = ('', 21)
        server = MultiprocessFTPServer(address, handler)

        # set a limit for connections
        server.max_cons = 256
        server.max_cons_per_ip = 5

        # start ftp server
        print("Task sheets are available over ftp on port 21\n" +
              "\tUsername: %s\n" % ("teacher" if self.teacher else "student") +
              "\tPassword: %s\n" % pw +
              "Keep this terminal open until you've retreived your files.")
        logging.basicConfig(level=logging.CRITICAL)
        server.serve_forever()

    def run_network(self):
        """Starts the Containernet network."""
        self.net.start()
