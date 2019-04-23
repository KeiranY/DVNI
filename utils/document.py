"""
Documenation
======================
"""


import os
import uuid
from docx import Document
from container import Docker
from utils import subnet
import networkx as nx
import matplotlib.pyplot as plt


def docker_hosts(net, doc=Document()):
    # If there are Docker instances
    # TODO: this could just find net.hosts that are instances of docker
    if Docker.added:
        doc.add_heading('Docker Hosts', level=1)
        for host in Docker.added:
            host.document(doc)
    return doc


def all_hosts(net, doc=Document()):
    # If there are Hosts
    if net.hosts:
        doc.add_heading('Address list', level=2)
        p = doc.add_paragraph()
        for host in net.hosts:
            if host.intfs:
                for key, intf in host.intfs.items():
                    p.add_run(('IP Address:\t%s/%s' % (str(intf.ip), str(intf.prefixLen))).ljust(30))
                    p.add_run('\tMAC:\t%s\n' % str(intf.mac))
    return doc


def subnet_table(net, doc=Document()):
    # If there are generated subnets
    if subnet.networks:
        doc.add_heading('Subnet Table', level=2)
        table = doc.add_table(rows=0, cols=3, style="Table Grid")
        header = table.add_row().cells
        header[0].text = "Subnet Address"
        header[1].text = "Subnet Mask"
        header[2].text = "Broadcast Address"
        for network in subnet.networks:
            row = table.add_row().cells
            row[0].text = str(network)
            row[1].text = str(network.netmask)
            row[2].text = str(network.broadcast_address)
    return doc


def switch_graph(net, doc=Document()):
    if net.switches:
        graph = nx.Graph()
        # Add all switches to the graph
        for switch in net.switches:
            graph.add_node(switch.name)
        # Get all links between 2 switches
        for link in [l for l in net.links if (l.intf1.node in net.switches and l.intf2.node in net.switches)]:
            graph.add_edge(link.intf1.node.name, link.intf2.node.name)
        nx.draw(graph, with_labels=True, node_size=1500, node_color="skyblue")
        pngName = str(uuid.uuid4()) + ".png"
        plt.savefig(pngName)
        doc.add_heading('Network Diagram', level=2)
        doc.add_picture(pngName)
        os.remove(pngName)
    return doc


def writeAnswers(net, doc=Document()):
    doc = docker_hosts(net, doc)
    doc = subnet_table(net, doc)
    doc = switch_graph(net, doc)
    doc = all_hosts(net, doc)
    return doc


def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.
    Source: https://github.com/python-openxml/python-docx/issues/74#issuecomment-261169410

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """
    import docx

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part

    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), "0000EE")
    rPr.append(c)

    # Add underline
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink
