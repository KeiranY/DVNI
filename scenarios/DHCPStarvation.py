"""
DHCP Starvation
=================
"""

# noinspection PyUnresolvedReferences
import mininet.node
from docx import Document
from mininet.link import TCLink

from scenarios import Scenario, DHCPIntro

# Use yersinia to send a RAW packet, is there a response? (Yes)
# What IP was the response sent from & to (Dhcp IP to broadcast)
# What are the network and netmask addresses for the response IP i.e. 192.168.0.1/255.255.255.0 (Dhcp ip and subnet netmask)
# What IP were you offered.

# Use yersinia to send DISCOVER packets, these will be sent continuously until you stop them
from utils import subnet


class Import(DHCPIntro.Import):
    """This scenario asks the user to perform a DHCP starvation attack and to report on the results."""

    name = "DHCP Starvation"
    enabled = True
    weight = 56

    def __init__(self, **kwargs):
        super(Import, self).__init__(**kwargs)

    def create_network(self, controller=None):
        """Extends the base scenario. Replaces the Kali containers link with a lower bandwidth one. This is so that the flooding attack dosn't cause a Denial of Service for the switch."""
        super(Import, self).create_network(controller)

        # Remove Kali's link and replace it with a bandwidth limited one, so our starvation attack doesnt topple the network
        link = self.kali.defaultIntf().link
        # Get the switch, the switch is whichever node in the link that kali isn't
        switch = link.intf1.node if link.intf2.node is self.kali else link.intf2.node
        self.net.removeLink(link)
        self.net.addLink(switch, self.kali, cls=TCLink, bw=0.25)

    def generate_task(self, doc=Document()):
        Scenario.generate_task(self, doc)
        doc.add_paragraph(
            "DHCP servers assign IP addresses within a set range for a set amount of time. "
            "This makes them vulnerable to 'starvation' attacks, where a malicious device requests every IP by pretending to be many devices simultaneously.")
        doc.add_paragraph(
            "In this task you'll be performing a starvation attack. To do this you'll need to use Yersinia. NOTE: Do NOT run dhclient until told.")
        doc.add_paragraph(
            "Answer the following questions:")

    def generate_questions(self):
        self.questions += [("Open wireshark/tcpdump to capture packets on the network and use Yersinia to send a RAW DHCP packet", ""),
                           ("What IP responded", self.dhcpd.IP()),
                           ("What IP did it respond to", "255.255.255.255"),
                           ("What are the network address and netmask for the network that responded", "%s & %s" % (str(subnet.networks[-1].network_address), str(subnet.networks[-1].netmask))),
                           ("How many IP addresses could be used in this network", str(len(list(subnet.networks[-1].hosts())))),
                           ("What IP was offered in the response", str(self.ips[0])),
                           ("Now use Yersinia to send DISCOVER packets, these will be sent continuously until you stop them in 'list attacks'. "
                            "After 10+ seconds check wireshark, filter for responses from the server offering IP addresses (bootp.option.dhcp==2). "
                            "How many offers were received", "Should be %s as all IP adresses in the network are unused excluding the DHCP server IP." % str(len(list(subnet.networks[-1].hosts()))-1)),
                           ("How long before all offers ceased, and the DHCP server was 'starved'", "~%.2f during testing, but any reasonable answer is acceptable." % (0.015*len(list(subnet.networks[-1].hosts()))-1)),
                           ("Run the starvation attack again, after a short wait attempt to use dhclient to get an IP address, does this work", "No")]


if __name__ == "__main__":
    # Run network
    Import(developer=True, seed="debug").run()
