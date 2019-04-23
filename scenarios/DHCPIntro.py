"""
DHCP Scanning
=============
"""
from docx import Document
from ipaddress import IPv4Address
from mininet.node import Controller, OVSSwitch, Host
from typing import List, Any

from container.dhcpd import Dhcpd, globals
from container.kali import Kali
from scenarios import Scenario
from utils import subnet


class Import(Scenario):
    """This scenario introduces the user to DHCP and it's relevance to IP addresses.
    It requires the use of packet sniffing and condiguration to find information about the network."""

    name = "DHCP Example"
    enabled = True
    weight = 55

    kali = None  # type: Kali
    dhcpd = None  # type: Dhcpd
    ips = None  # type: List[IPv4Address]
    network = None

    def create_network(self, controller=None):
        """Extends the base class. Adds a DHCP server and a Kali container."""
        super(Import, self).create_network(controller)

        prefix_length = 27
        self.network = subnet.generate(prefix_length)
        self.ips = list(self.network.hosts())

        switch = self.net.addSwitch('s1', cls=OVSSwitch, failMode='standalone')
        self.dhcpd = self.net.addDocker('dhcpd',
                                        cls=Dhcpd,
                                        ip="%s/%s" % (self.ips.pop(0), prefix_length))  # type: Dhcpd
        self.dhcpd.add_subnet(self.network, [globals.default_lease_time(120),
                                             globals.min_lease_time(60),
                                             globals.max_lease_time(120)])
        self.kali = self.net.addDocker('kali',
                                       cls=Kali,
                                       ip="0.0.0.0/32")

        self.kali.install_package('yersinia', 'wireshark')

        self.net.addLink(switch, self.dhcpd)
        self.net.addLink(switch, self.kali)

    def generate_task(self, doc=Document()):
        super(Import, self).generate_task(doc)
        doc.add_paragraph(
            "In this task your Kali machine is not assigned an IP on the network. There is a DHCP server able to assign you one.")
        doc.add_paragraph(
            "Answer the following questions:")

    def generate_questions(self):
        self.questions = \
            [("What is the IP of the %s adapter", str(self.kali.defaultIntf())),
             ("In a terminal use dhclient on the %s adapter, what is the IP now",
              "Any IP in the range %s to %s is valid" % (str(self.network.network_address + 1), str(self.network.broadcast_address - 1))),
             ("Using wireshark/tcpdump find an DHCP ACK or OFFER packet and find the following information within it:\n\tServer identifier/ip", str(self.dhcpd.IP())),
             ("\tIP Lease time", "60")]


if __name__ == "__main__":
    # Run network
    Import(developer=True, seed="debug").run()
